#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Advanced Markdown to Word Document Converter for Course Materials
Handles complex formatting requirements including colored text blocks, timecodes, etc.
"""

import re
from pathlib import Path
from datetime import datetime
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

class CourseDocConverter:
    def __init__(self):
        self.doc = None
        self.current_unit = 0
        self.unit_start_positions = {}

    def setup_document(self):
        """Setup document with proper formatting"""
        self.doc = Document()
        section = self.doc.sections[0]
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(2.54)
        section.right_margin = Cm(2.54)

        # Set footer with course name and page number
        footer = section.footer
        footer_para = footer.paragraphs[0]
        footer_para.text = "頁碼：第 "
        run = footer_para.add_run()
        fldChar1 = OxmlElement('w:fldChar')
        fldChar1.set(qn('w:fldCharType'), 'begin')
        run._r.append(fldChar1)

        instrText = OxmlElement('w:instrText')
        instrText.set(qn('xml:space'), 'preserve')
        instrText.text = 'PAGE'
        run._r.append(instrText)

        fldChar2 = OxmlElement('w:fldChar')
        fldChar2.set(qn('w:fldCharType'), 'end')
        run._r.append(fldChar2)

        footer_para.add_run(" 頁")

        # Setup styles
        self._setup_styles()

    def _setup_styles(self):
        """Setup custom paragraph styles"""
        style = self.doc.styles['Normal']
        style.font.name = '微軟雅黑'
        style.font.size = Pt(11)
        style.paragraph_format.line_spacing = 1.5

        for i in range(1, 5):
            try:
                heading = self.doc.styles[f'Heading {i}']
                heading.font.name = '微軟雅黑'
                heading.paragraph_format.line_spacing = 1.5
            except:
                pass

    def add_page_break(self):
        """Add page break"""
        self.doc.add_page_break()

    def add_colored_box(self, text, color):
        """Add colored text box"""
        p = self.doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.5)
        p.paragraph_format.right_indent = Cm(0.5)
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(6)

        run = p.add_run(text)

        # Add background color
        shading = OxmlElement('w:shd')
        shading.set(qn('w:fill'), color)
        p._element.get_or_add_pPr().append(shading)

        return p

    def add_cover_page(self, course_name, duration, audience):
        """Create professional cover page"""
        # Spacing
        for _ in range(4):
            self.doc.add_paragraph()

        # Title
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(course_name)
        run.font.size = Pt(32)
        run.bold = True
        run.font.color.rgb = RGBColor(0, 51, 102)

        # Spacing
        for _ in range(3):
            self.doc.add_paragraph()

        # Course info
        info = [
            ('課程時長', duration),
            ('學習受眾', audience),
            ('授課日期', '__________________________'),
            ('講師名字', '__________________________')
        ]

        for label, value in info:
            p = self.doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(10)
            p.paragraph_format.space_after = Pt(10)
            run = p.add_run(f'{label}：{value}')
            run.font.size = Pt(12)

        # Footer
        for _ in range(5):
            self.doc.add_paragraph()

        p = self.doc.add_paragraph(f'文檔生成日期：{datetime.now().strftime("%Y 年 %m 月 %d 日")}')
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        self.add_page_break()

    def add_prep_section(self):
        """Add course preparation section"""
        # Title
        p = self.doc.add_paragraph()
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(6)
        run = p.add_run('課程準備清單')
        run.font.size = Pt(16)
        run.bold = True

        self.doc.add_paragraph('請在授課前 24 小時完成以下檢查項目：', style='Normal')

        checklist = [
            '【技術檢查】投影設備、電腦、網路連接、ChatGPT 訪問權限驗證',
            '【教材準備】講義列印、案例文件準備、備用筆記本和文具',
            '【學員準備】確認學員人數、發放上課帳號、準備簽到表',
            '【現場佈置】教室溫度、飲水及茶點、座位安排（電腦教室）',
            '【演練測試】完整走過所有演示案例和互動環節',
            '【應急預案】網路中斷時的備用方案、時間超支的應對'
        ]

        for item in checklist:
            self.doc.add_paragraph(item, style='List Bullet')

        self.add_page_break()

    def add_progress_table(self, units_list):
        """Add course progress and timeline table"""
        p = self.doc.add_paragraph()
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(6)
        run = p.add_run('課程進度表 & 時間分配')
        run.font.size = Pt(14)
        run.bold = True

        table = self.doc.add_table(rows=1, cols=5)
        table.style = 'Light Grid Accent 1'

        # Header
        headers = ['單元序號', '單元標題', '預估時長', '講述 vs 討論', '關鍵活動']
        for i, header in enumerate(headers):
            cell = table.rows[0].cells[i]
            cell.text = header
            for para in cell.paragraphs:
                for run in para.runs:
                    run.bold = True

        # Add unit rows
        for idx, unit_title in enumerate(units_list, 1):
            row = table.add_row()
            row.cells[0].text = f'Unit {idx}'
            row.cells[1].text = unit_title
            row.cells[2].text = '35-45 min'
            row.cells[3].text = '見課程內容'
            row.cells[4].text = '見課程內容'

        self.doc.add_paragraph()
        self.add_page_break()

    def process_markdown(self, content, course_name):
        """Process markdown content and add to document"""
        lines = content.split('\n')

        # Extract course info
        duration_match = re.search(r'\*\*總時長\*\*：(.+)', content)
        duration = duration_match.group(1).strip() if duration_match else '3.5 小時'

        audience_match = re.search(r'\*\*受眾\*\*：(.+)', content)
        audience = audience_match.group(1).strip() if audience_match else ''

        # Add cover and prep
        self.add_cover_page(course_name, duration, audience)
        self.add_prep_section()

        # Extract units for progress table
        units = re.findall(r'## (第.{1,2}單元：.+)', content)
        if units:
            clean_units = [u.replace('## ', '').replace('第', '').replace('單元：', '')
                          for u in units[:6]]
            self.add_progress_table(clean_units)

        # Process main content
        self._process_content_lines(lines)

    def _process_content_lines(self, lines):
        """Process content lines and add appropriate formatting"""
        i = 0
        skip_toc = True
        in_code_block = False
        unit_count = 0

        while i < len(lines):
            line = lines[i]

            # Skip TOC section
            if skip_toc:
                if line.startswith('---') and i > 10:
                    skip_toc = False
                    i += 1
                    continue
                i += 1
                continue

            # Handle code blocks
            if line.startswith('```'):
                in_code_block = not in_code_block
                i += 1
                continue

            if in_code_block:
                p = self.doc.add_paragraph(line, style='Normal')
                for run in p.runs:
                    run.font.name = 'Courier New'
                    run.font.size = Pt(9)
                i += 1
                continue

            # Process headings
            if line.startswith('## 第') and '單元' in line:
                unit_count += 1
                if unit_count > 1:
                    self.add_page_break()

                # Extract unit title
                unit_title = line.replace('## ', '').strip()
                p = self.doc.add_paragraph(unit_title, style='Heading 1')
                p.paragraph_format.space_before = Pt(18)
                p.paragraph_format.space_after = Pt(12)

                i += 1
                continue

            elif line.startswith('## '):
                title = line.replace('## ', '').strip()
                p = self.doc.add_paragraph(title, style='Heading 1')
                p.paragraph_format.space_before = Pt(12)
                p.paragraph_format.space_after = Pt(6)
                i += 1
                continue

            elif line.startswith('### '):
                title = line.replace('### ', '').strip()
                p = self.doc.add_paragraph(title, style='Heading 2')
                p.paragraph_format.space_before = Pt(10)
                p.paragraph_format.space_after = Pt(4)
                i += 1
                continue

            elif line.startswith('#### '):
                title = line.replace('#### ', '').strip()
                p = self.doc.add_paragraph(title, style='Heading 3')
                p.paragraph_format.space_before = Pt(8)
                p.paragraph_format.space_after = Pt(2)
                i += 1
                continue

            # Handle timecodes
            elif re.match(r'^\*\d+:\d+', line):
                self._add_timecode_para(line)
                i += 1
                continue

            # Handle instructor notes (gray box)
            elif '講稿（講師使用）' in line or line.strip().startswith('講稿'):
                # Consume next lines as instructor content
                i += 1
                instructor_text = []
                while i < len(lines) and lines[i].strip() and not lines[i].startswith('#'):
                    if lines[i].startswith('####'):
                        break
                    instructor_text.append(lines[i])
                    i += 1

                if instructor_text:
                    content = '\n'.join(instructor_text).strip()
                    self.add_colored_box(content, 'E8E8E8')
                continue

            # Handle activity markers (yellow box)
            elif '活動指引' in line or '互動' in line or '【活動】' in line:
                self.add_colored_box('【互動活動時間】', 'FFFACD')
                i += 1
                continue

            # Handle Q&A markers (light blue)
            elif 'Q&A' in line or '問答' in line:
                self.add_colored_box('【學員提問和講師回答時間】', 'ADD8E6')
                i += 1
                continue

            # Handle bullet points
            elif line.startswith('- '):
                text = line[2:].strip()
                self.doc.add_paragraph(text, style='List Bullet')
                i += 1
                continue

            # Skip empty lines and markdown separators
            elif not line.strip() or line.startswith('---') or line.startswith('|'):
                i += 1
                continue

            # Regular paragraph
            elif line.strip():
                p = self.doc.add_paragraph(line)
                p.paragraph_format.line_spacing = 1.5

                # Apply formatting for bold and italic
                self._apply_inline_formatting(p)

                i += 1
                continue

            i += 1

    def _add_timecode_para(self, text):
        """Add timecode paragraph with monospace font"""
        p = self.doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(1)
        p.paragraph_format.space_before = Pt(3)
        p.paragraph_format.space_after = Pt(3)

        run = p.add_run(text)
        run.font.name = 'Courier New'
        run.font.size = Pt(10)

    def _apply_inline_formatting(self, paragraph):
        """Apply inline formatting (bold, italic) to paragraph"""
        # Simple regex-based formatting
        text = paragraph.text

        # Bold patterns **text** or __text__
        if '**' in text or '__' in text:
            paragraph.clear()
            parts = re.split(r'\*\*(.+?)\*\*|__(.+?)__', text)
            for i, part in enumerate(parts):
                if part:
                    run = paragraph.add_run(part)
                    if i % 2 == 1:  # Odd indices are bold
                        run.bold = True

        # Italic patterns *text* or _text_
        elif '*' in text:
            paragraph.clear()
            parts = re.split(r'\*(.+?)\*', text)
            for i, part in enumerate(parts):
                if part:
                    run = paragraph.add_run(part)
                    if i % 2 == 1:  # Odd indices are italic
                        run.italic = True

    def save(self, filepath):
        """Save document"""
        self.doc.save(filepath)

def convert_course_to_docx(md_filepath, docx_filepath, course_name):
    """Main conversion function"""
    try:
        # Read markdown
        with open(md_filepath, 'r', encoding='utf-8') as f:
            content = f.read()

        # Create converter
        converter = CourseDocConverter()
        converter.setup_document()
        converter.process_markdown(content, course_name)
        converter.save(docx_filepath)

        print(f'✓ 已生成：{docx_filepath}')
        return True
    except Exception as e:
        print(f'✗ 轉換失敗 {md_filepath}：{e}')
        import traceback
        traceback.print_exc()
        return False

def main():
    """Main entry point"""
    base_dir = Path('\\\\192.168.168.207\\tuq-0001\\projects\\goose\\AI training')

    print('開始轉換課程內容為 Word 文檔...')
    print()

    courses = [
        ('課程內容_AI實戰基礎入門.md', 'AI實戰基礎入門_講師版.docx', 'AI 實戰基礎入門'),
        ('課程內容_AI虛擬組織.md', 'AI虛擬組織_講師版.docx', 'AI 虛擬組織')
    ]

    success_count = 0
    for md_file, docx_file, course_name in courses:
        md_path = base_dir / md_file
        docx_path = base_dir / docx_file

        if convert_course_to_docx(str(md_path), str(docx_path), course_name):
            success_count += 1

    print()
    if success_count == len(courses):
        print(f'✓ 成功生成 {success_count} 份文檔！')
        print()
        print('輸出位置：')
        for _, docx_file, _ in courses:
            print(f'  • {base_dir / docx_file}')
    else:
        print(f'⚠ 僅成功生成 {success_count}/{len(courses)} 份文檔')

if __name__ == '__main__':
    main()

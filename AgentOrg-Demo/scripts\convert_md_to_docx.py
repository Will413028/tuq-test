#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Convert Markdown course content to professional Word documents
"""

import re
from pathlib import Path
from datetime import datetime
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def add_page_break(doc):
    """Add page break to document"""
    doc.add_page_break()

def set_cell_background(cell, color):
    """Set cell background color"""
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), color)
    cell._element.get_or_add_tcPr().append(shading_elm)

def add_colored_paragraph(doc, text, bg_color=None, style='Normal', bold=False, italic=False):
    """Add a paragraph with optional background color"""
    p = doc.add_paragraph()
    p.style = style
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic

    if bg_color:
        # Add background color to paragraph
        shading = OxmlElement('w:shd')
        shading.set(qn('w:fill'), bg_color)
        p._element.get_or_add_pPr().append(shading)

    return p

def add_instructor_note(doc, text):
    """Add instructor note with gray background"""
    add_colored_paragraph(doc, text, bg_color='D3D3D3')

def add_activity_note(doc, text):
    """Add activity note with light yellow background"""
    add_colored_paragraph(doc, text, bg_color='FFFACD')

def add_qa_note(doc, text):
    """Add Q&A note with light blue background"""
    add_colored_paragraph(doc, text, bg_color='ADD8E6')

def add_timecode(doc, text):
    """Add timecode with monospace font"""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = 'Courier New'
    run.font.size = Pt(10)
    return p

def parse_markdown_file(filepath):
    """Parse markdown file and extract sections"""
    with open(filepath, 'utf-8', encoding='utf-8') as f:
        content = f.read()
    return content

def create_cover_page(doc, course_name, duration, audience):
    """Create cover page"""
    # Add empty space
    for _ in range(3):
        doc.add_paragraph()

    # Add course name
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(course_name)
    run.font.size = Pt(28)
    run.bold = True

    # Add spacing
    for _ in range(2):
        doc.add_paragraph()

    # Add course info
    info_items = [
        f"課程時長：{duration}",
        f"受眾：{audience}",
        f"日期：_____________________（講師填寫）",
        f"講師：_____________________"
    ]

    for item in info_items:
        p = doc.add_paragraph(item)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(12)

    # Add footer
    for _ in range(4):
        doc.add_paragraph()

    p = doc.add_paragraph(f"生成時間：{datetime.now().strftime('%Y-%m-%d')}")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    add_page_break(doc)

def create_toc_page(doc):
    """Create table of contents placeholder"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("目錄")
    run.font.size = Pt(18)
    run.bold = True

    doc.add_paragraph("（此頁目錄將由 Word 自動生成，請使用「參考資料」→「目錄」功能更新）")

    add_page_break(doc)

def create_progress_table(doc):
    """Create course progress table"""
    p = doc.add_paragraph()
    run = p.add_run("課程進度表與時間分配")
    run.font.size = Pt(14)
    run.bold = True

    table = doc.add_table(rows=1, cols=4)
    table.style = 'Light Grid Accent 1'

    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = '單元'
    hdr_cells[1].text = '時長'
    hdr_cells[2].text = '講述時間'
    hdr_cells[3].text = '重點活動'

    # Make header bold
    for cell in hdr_cells:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True

    return table

def setup_document_styles(doc):
    """Setup document styles and formatting"""
    # Set default font for all paragraphs
    style = doc.styles['Normal']
    style.font.name = '微軟雅黑'
    style.font.size = Pt(11)
    style.paragraph_format.line_spacing = 1.5

    # Setup heading styles
    for i in range(1, 5):
        heading_style = doc.styles[f'Heading {i}']
        heading_style.font.name = '微軟雅黑'
        heading_style.paragraph_format.line_spacing = 1.5

def convert_markdown_to_docx(md_filepath, docx_filepath, is_basic=True):
    """Convert markdown to Word document"""

    # Read markdown content
    with open(md_filepath, 'r', encoding='utf-8') as f:
        content = f.read()

    # Create document
    doc = Document()

    # Setup document
    section = doc.sections[0]
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(2.54)
    section.right_margin = Cm(2.54)

    setup_document_styles(doc)

    # Extract course info from markdown
    course_name_match = re.search(r'\*\*課程名稱\*\*：(.+)', content)
    course_name = course_name_match.group(1).strip() if course_name_match else "課程"

    duration_match = re.search(r'\*\*總時長\*\*：(.+)', content)
    duration = duration_match.group(1).strip() if duration_match else ""

    audience_match = re.search(r'\*\*受眾\*\*：(.+)', content)
    audience = audience_match.group(1).strip() if audience_match else ""

    # Create cover page
    create_cover_page(doc, course_name, duration, audience)

    # Create progress page
    p = doc.add_paragraph()
    run = p.add_run("課程準備清單")
    run.font.size = Pt(16)
    run.bold = True

    doc.add_paragraph("請在授課前確認以下事項：")

    checklist = [
        "投影設備可正常運作，測試分辨率和音量",
        "學員電腦已安裝所需工具和軟體",
        "網絡連接穩定，ChatGPT 或等價服務可訪問",
        "準備紙質筆記本和筆供學員使用",
        "預留飲用水和簡單點心",
        "測試所有範例代碼/案例",
        "確認每位學員收到上課資料"
    ]

    for item in checklist:
        doc.add_paragraph(item, style='List Bullet')

    add_page_break(doc)

    # Create course progress table
    table = create_progress_table(doc)

    # Add sample rows based on course structure
    units = re.findall(r'^## 第[一二三四五六七八九]單元：(.+)', content, re.MULTILINE)
    for idx, unit_name in enumerate(units[:5], 1):
        row = table.add_row()
        row.cells[0].text = f"第 {idx} 單元"
        row.cells[1].text = f"{45 - idx*5} 分鐘"
        row.cells[2].text = "見下方詳細內容"
        row.cells[3].text = "見下方詳細內容"

    doc.add_paragraph()
    add_page_break(doc)

    # Parse and add content
    lines = content.split('\n')
    current_heading = None
    skip_toc = True

    for i, line in enumerate(lines):
        # Skip course info section at beginning
        if skip_toc and (line.startswith('## 課程資訊') or line.startswith('- **')):
            continue
        if skip_toc and line.startswith('---'):
            skip_toc = False
            continue

        # Skip empty lines at start of parsing
        if not line.strip() and current_heading is None:
            continue

        # Process headings
        if line.startswith('# '):
            continue
        elif line.startswith('## '):
            current_heading = line[3:].strip()
            p = doc.add_paragraph(current_heading, style='Heading 1')
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after = Pt(6)
        elif line.startswith('### '):
            heading = line[4:].strip()
            p = doc.add_paragraph(heading, style='Heading 2')
            p.paragraph_format.space_before = Pt(10)
            p.paragraph_format.space_after = Pt(4)
        elif line.startswith('#### '):
            heading = line[5:].strip()
            p = doc.add_paragraph(heading, style='Heading 3')
            p.paragraph_format.space_before = Pt(8)
            p.paragraph_format.space_after = Pt(2)

        # Handle special content markers
        elif '講稿（講師使用）' in line or '【講稿】' in line:
            add_instructor_note(doc, "【講師講稿區】")
        elif '【活動指引】' in line or '活動指引' in line:
            add_activity_note(doc, "【互動活動區】")
        elif '【Q&A】' in line or 'Q&A' in line:
            add_qa_note(doc, "【學員 Q&A 時間】")
        elif re.match(r'^\*\d+:\d+', line):
            # Timecode line
            add_timecode(doc, line)
        elif line.startswith('- '):
            # Bullet point
            doc.add_paragraph(line[2:], style='List Bullet')
        elif line.startswith('| '):
            # Skip markdown tables (we'll handle them separately)
            continue
        elif line.strip() and not line.startswith(('---', '```', '*', '#')):
            # Regular paragraph
            if line.startswith('*') and line.endswith('*'):
                # Italic text
                p = doc.add_paragraph()
                run = p.add_run(line.strip('*'))
                run.italic = True
            else:
                p = doc.add_paragraph(line)
                p.paragraph_format.line_spacing = 1.5

        # Add page break after major sections
        if current_heading and '單元' in current_heading and i > 0:
            # Check if we're at the start of a new unit
            if i < len(lines) - 1 and re.match(r'^## 第.單元', lines[i+1]):
                add_page_break(doc)

    # Save document
    doc.save(docx_filepath)
    print(f"✓ 已生成：{docx_filepath}")

def main():
    """Main conversion process"""

    base_dir = Path('\\\\192.168.168.207\\tuq-0001\\projects\\goose\\AI training')

    # Course 1: AI 實戰基礎入門
    md_file1 = base_dir / '課程內容_AI實戰基礎入門.md'
    docx_file1 = base_dir / 'AI實戰基礎入門_講師版.docx'

    # Course 2: AI 虛擬組織
    md_file2 = base_dir / '課程內容_AI虛擬組織.md'
    docx_file2 = base_dir / 'AI虛擬組織_講師版.docx'

    print("開始轉換課程內容為 Word 文檔...")
    print()

    try:
        convert_markdown_to_docx(str(md_file1), str(docx_file1), is_basic=True)
        convert_markdown_to_docx(str(md_file2), str(docx_file2), is_basic=False)
        print()
        print("✓ 所有文檔已成功生成！")
        print()
        print("輸出位置：")
        print(f"  1. {docx_file1}")
        print(f"  2. {docx_file2}")
    except Exception as e:
        print(f"✗ 轉換過程中出錯：{e}")
        import traceback
        traceback.print_exc()

if __name__ == '__main__':
    main()
